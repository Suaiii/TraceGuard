"""Official competition report DOCX builder tests."""

from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from scripts.build_report_docx import build_report, split_report_sections


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = PROJECT_ROOT / "作品报告模板.docx"
requires_official_template = pytest.mark.skipif(
    not TEMPLATE_PATH.is_file(),
    reason="official report template is an external competition asset",
)


def test_split_report_sections_finds_required_sections():
    sections = split_report_sections(PROJECT_ROOT / "reports" / "TraceGuard.md")

    assert list(sections) == ["摘要", "第一章", "第二章", "第三章", "第四章", "第五章", "参考文献"]
    assert "Facebook" in sections["第三章"]


@requires_official_template
def test_build_report_preserves_template_sections_and_inserts_evidence(tmp_path):
    output = tmp_path / "TraceGuard_report.docx"

    build_report(
        TEMPLATE_PATH,
        PROJECT_ROOT / "reports" / "TraceGuard.md",
        output,
        project_root=PROJECT_ROOT,
    )

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert len(document.sections) == 10
    assert "TraceGuard：面向跨域 AIGC 图像的可解释伪造检测与篡改取证平台" in text
    assert "本部分内容主要说明作品的创新性" not in text
    assert len(document.inline_shapes) >= 6
    assert all(
        shape._inline.docPr.get("descr")
        for shape in document.inline_shapes
    )
    assert len(document.tables) >= 4
    with ZipFile(output) as package:
        footer_xml = package.read("word/footer2.xml").decode("utf-8")
    assert "NUMPAGES" in footer_xml
    assert "- 3" in footer_xml
