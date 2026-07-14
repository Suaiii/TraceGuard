"""Build the official TraceGuard competition report from the Markdown source."""

import argparse
import re
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.parts.numbering import NumberingPart
from docx.shared import Inches, Pt


# Office's built-in MathML -> OMML transform, used to emit natively editable
# Word equations instead of rasterised LaTeX images.
MML2OMML_XSL = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
_MML2OMML_TRANSFORM = None


def _get_mml2omml_transform():
    """Return a cached lxml XSLT transform for MathML -> OMML."""
    global _MML2OMML_TRANSFORM
    if _MML2OMML_TRANSFORM is None:
        from lxml import etree

        _MML2OMML_TRANSFORM = etree.XSLT(etree.parse(MML2OMML_XSL))
    return _MML2OMML_TRANSFORM


def _latex_to_omath_xml(expression):
    """Convert a LaTeX expression to an OMML ``<m:oMath>`` XML byte string."""
    import latex2mathml.converter
    from lxml import etree

    latex = expression.strip()
    if latex.startswith("$$") and latex.endswith("$$"):
        latex = latex[2:-2]
    else:
        latex = latex.strip("$")
    latex = latex.strip()
    mathml = latex2mathml.converter.convert(latex)
    mathml_element = etree.fromstring(mathml.encode("utf-8"))
    omml_root = _get_mml2omml_transform()(mathml_element).getroot()
    return etree.tostring(omml_root)


# Inline math handling. A redundant LaTeX ``\( ... \)`` wrapper around an inline
# ``$...$`` span is a dirty authoring artefact and is unwrapped first; genuine
# bare parentheses in prose are left untouched. Extracted spans are swapped for
# private-use sentinels so ``_strip_inline`` can run on the surrounding prose
# without ever mangling the LaTeX (e.g. its ``\_`` / ``\(`` sequences).
_INLINE_MATH_WRAP_RE = re.compile(r"\\\(\s*(\$[^$]+\$)\s*\\\)")
_INLINE_MATH_RE = re.compile(r"\$([^$]+)\$")
_MATH_SENTINEL_RE = re.compile("(\\d+)")


def _extract_inline_math(raw_text):
    """Replace inline ``$...$`` spans with sentinels; return (text, [latex])."""
    text = _INLINE_MATH_WRAP_RE.sub(lambda m: m.group(1), raw_text)
    exprs = []

    def _swap(match):
        exprs.append(match.group(1))
        return "%d" % (len(exprs) - 1)

    return _INLINE_MATH_RE.sub(_swap, text), exprs


def _add_inline_math(paragraph, latex):
    """Append an inline ``<m:oMath>`` to the paragraph; True on success."""
    try:
        paragraph._p.append(parse_xml(_latex_to_omath_xml(latex)))
        return True
    except Exception:
        return False


def _add_inline_content(paragraph, raw_text):
    """Populate a paragraph from raw markdown, rendering inline ``$...$`` as OMML.

    Inline math is pulled from the *raw* text before ``_strip_inline`` touches it;
    the remaining prose still flows through ``_strip_inline`` exactly as before.
    A formula that fails to convert falls back to its literal ``$...$`` text.
    """
    placeholder_text, exprs = _extract_inline_math(raw_text)
    stripped = _strip_inline(placeholder_text)
    pos = 0
    for match in _MATH_SENTINEL_RE.finditer(stripped):
        if match.start() > pos:
            paragraph.add_run(stripped[pos:match.start()])
        latex = exprs[int(match.group(1))]
        if not _add_inline_math(paragraph, latex):
            paragraph.add_run("$%s$" % latex)
        pos = match.end()
    if pos < len(stripped) or not paragraph.runs and pos == 0:
        paragraph.add_run(stripped[pos:])
    return paragraph


SECTION_HEADINGS = OrderedDict([
    ("摘要", "摘要"),
    ("第一章", "第一章 作品概述"),
    ("第二章", "第二章 作品设计与实现"),
    ("第三章", "第三章 作品测试与分析"),
    ("第四章", "第四章 创新性说明"),
    ("第五章", "第五章 总结"),
    ("参考文献", "参考文献"),
])
PROJECT_TITLE = "TraceGuard：面向跨域 AIGC 图像的可解释伪造检测与篡改取证平台"


def split_report_sections(path):
    text = Path(path).read_text(encoding="utf-8")
    pattern = re.compile(
        r"(?ms)^# (摘要|第一章[^\r\n]*|第二章[^\r\n]*|第三章[^\r\n]*|"
        r"第四章[^\r\n]*|第五章[^\r\n]*|参考文献)\s*(.*?)(?=^# |\Z)"
    )
    found = {}
    for match in pattern.finditer(text):
        heading = match.group(1)
        key = next(key for key in SECTION_HEADINGS if heading.startswith(key))
        found[key] = match.group(2).strip()
    missing = [key for key in SECTION_HEADINGS if key not in found]
    if missing:
        raise ValueError(f"missing report sections: {missing}")
    return OrderedDict((key, found[key]) for key in SECTION_HEADINGS)


def _strip_inline(text):
    text = re.sub(r"!\[[^]]*\]\([^)]+\)", "", text)
    text = text.replace("**", "").replace("`", "")
    text = text.replace("\\.", ".").replace("\\-", "-")
    text = text.replace("\\[", "[").replace("\\]", "]")
    text = text.replace("\\(", "(").replace("\\)", ")")
    text = text.replace("\\_", "_")
    text = re.sub(r"(?<!\*)\*([^*]+)\*", r"\1", text)
    return text.strip()


def _parse_table(lines):
    rows = []
    for line in lines:
        cells = [_strip_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows


def parse_blocks(markdown):
    lines = markdown.splitlines()
    blocks = []
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        image = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", line)
        if image:
            blocks.append(("image", image.group(1), image.group(2)))
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            blocks.append(("table", _parse_table(table_lines)))
            continue
        if line.startswith("$") and line.endswith("$"):
            blocks.append(("equation", line))
            index += 1
            continue
        if re.match(r"^\d+\\?\.\d+", line):
            blocks.append(("heading2", line))
            index += 1
            continue
        if line.startswith("**") and line.endswith("**"):
            blocks.append(("heading3", line))
            index += 1
            continue
        if line.startswith("- "):
            blocks.append(("bullet", line[2:]))
            index += 1
            continue
        if re.match(r"^(图|表)\s*\d", line):
            blocks.append(("caption", line))
            index += 1
            continue
        blocks.append(("paragraph", line))
        index += 1
    return blocks


def _set_run_font(run, name, size, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _replace_paragraph(paragraph, text, name="黑体", size=16, bold=True):
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    _set_run_font(run, name, size, bold=bold)
    return paragraph


def _append_field_run(paragraph, field_type=None, instruction=None, text=None):
    run = OxmlElement("w:r")
    if field_type is not None:
        field = OxmlElement("w:fldChar")
        field.set(qn("w:fldCharType"), field_type)
        run.append(field)
    elif instruction is not None:
        node = OxmlElement("w:instrText")
        node.set(qn("xml:space"), "preserve")
        node.text = instruction
        run.append(node)
    elif text is not None:
        node = OxmlElement("w:t")
        node.text = text
        run.append(node)
    paragraph._p.append(run)


def _replace_static_total_page_count(document):
    seen_parts = set()
    for section in document.sections:
        footer = section.footer
        if id(footer.part) in seen_parts:
            continue
        seen_parts.add(id(footer.part))
        for paragraph in footer.paragraphs:
            if "共" not in paragraph.text:
                continue
            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)
            text_run = paragraph.add_run("第 ")
            _set_run_font(text_run, "宋体", 10.5)
            _append_field_run(paragraph, field_type="begin")
            _append_field_run(paragraph, instruction=" PAGE ")
            _append_field_run(paragraph, field_type="separate")
            _append_field_run(paragraph, text="1")
            _append_field_run(paragraph, field_type="end")
            text_run = paragraph.add_run(" 页 共 ")
            _set_run_font(text_run, "宋体", 10.5)
            _append_field_run(paragraph, field_type="begin")
            _append_field_run(paragraph, instruction=" = ")
            _append_field_run(paragraph, field_type="begin")
            _append_field_run(paragraph, instruction=" NUMPAGES ")
            _append_field_run(paragraph, field_type="separate")
            _append_field_run(paragraph, text="10")
            _append_field_run(paragraph, field_type="end")
            _append_field_run(paragraph, instruction=" - 3 ")
            _append_field_run(paragraph, field_type="separate")
            _append_field_run(paragraph, text="7")
            _append_field_run(paragraph, field_type="end")
            text_run = paragraph.add_run("页")
            _set_run_font(text_run, "宋体", 10.5)


def _format_body(paragraph, first_line=True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(24) if first_line else Pt(0)
    for run in paragraph.runs:
        _set_run_font(run, "宋体", 12)


def _format_heading(paragraph, level=2):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(6 if level == 2 else 3)
    fmt.space_after = Pt(0)
    fmt.keep_with_next = True
    fmt.first_line_indent = Pt(0)
    for run in paragraph.runs:
        _set_run_font(run, "黑体", 12, bold=True)


def _shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _configure_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), "8712")
    tbl_width.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "808080")
        borders.append(border)
    tbl_pr.append(borders)
    column_count = len(table.columns)
    width = Inches(6.05 / column_count)
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for cell in row.cells:
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if row_index == 0:
                _shade_cell(cell, "E7ECEA")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_run_font(run, "宋体", 8.5, bold=row_index == 0)


def _ensure_bullet_numbering(document):
    cached = getattr(document, "_traceguard_bullet_num_id", None)
    if cached is not None:
        return cached
    try:
        numbering_part = document.part.part_related_by(RT.NUMBERING)
    except KeyError:
        numbering_xml = (
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            CT.WML_NUMBERING,
            parse_xml(numbering_xml),
            document.part.package,
        )
        document.part.relate_to(numbering_part, RT.NUMBERING)
    numbering = numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    document._traceguard_bullet_num_id = num_id
    return num_id


def _apply_bullet(document, paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(_ensure_bullet_numbering(document)))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)


def _section_anchors(document):
    headings = {}
    for paragraph in document.paragraphs:
        normalized = paragraph.text.replace(" ", "")
        for key, expected in SECTION_HEADINGS.items():
            if normalized == expected.replace(" ", ""):
                headings[key] = paragraph._p
    if set(headings) != set(SECTION_HEADINGS):
        raise ValueError(f"template headings not found: {set(SECTION_HEADINGS) - set(headings)}")

    body = document._element.body
    children = list(body)
    anchors = {}
    for position, key in enumerate(SECTION_HEADINGS):
        heading = headings[key]
        start = children.index(heading)
        if position + 1 < len(SECTION_HEADINGS):
            next_heading = headings[list(SECTION_HEADINGS)[position + 1]]
            stop = children.index(next_heading)
            candidates = children[start + 1:stop]
            section_break = next(
                (node for node in reversed(candidates) if node.find(".//" + qn("w:sectPr")) is not None),
                next_heading,
            )
            anchors[key] = section_break
        else:
            anchors[key] = body.sectPr
    return headings, anchors


def _clear_section_content(document, heading, anchor):
    body = document._element.body
    children = list(body)
    start = children.index(heading)
    stop = children.index(anchor)
    for node in children[start + 1:stop]:
        body.remove(node)


def _move_before(anchor, element):
    anchor.addprevious(element)


def _add_paragraph_before(document, anchor, text, kind):
    paragraph = document.add_paragraph()
    # ``text`` is now the raw markdown line: render inline $...$ as OMML while the
    # surrounding prose still flows through _strip_inline. The _format_* helpers
    # below re-apply fonts to every text run (inline oMath elements are untouched).
    _add_inline_content(paragraph, text)
    if kind in {"heading2", "heading3"}:
        _format_heading(paragraph, 2 if kind == "heading2" else 3)
    elif kind == "caption":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.keep_with_next = text.startswith("表")
        for run in paragraph.runs:
            _set_run_font(run, "宋体", 10.5)
    else:
        _format_body(paragraph, first_line=kind == "paragraph")
        if kind == "bullet":
            _apply_bullet(document, paragraph)
    _move_before(anchor, paragraph._p)
    return paragraph


def _add_image_before(document, anchor, path, width=5.8, alt_text=""):
    path = Path(path)
    if path.suffix.lower() == ".svg":
        raster_path = path.with_suffix(".png")
        if not raster_path.is_file():
            raise FileNotFoundError(f"PNG fallback not found for SVG: {path}")
        path = raster_path
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    inline_shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    description = alt_text.strip() or path.stem.replace("_", " ")
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", description)
    _move_before(anchor, paragraph._p)


def _add_equation_before(document, anchor, expression, output_dir, index):
    # Primary path: emit a natively editable Word equation (OMML) built from the
    # LaTeX source via latex2mathml -> Office MML2OMML.XSL. Any failure falls
    # back to the matplotlib PNG rendering below so the build never crashes.
    try:
        omath_xml = _latex_to_omath_xml(expression)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph._p.append(parse_xml(omath_xml))
        _move_before(anchor, paragraph._p)
        return
    except Exception:
        pass

    from matplotlib.mathtext import math_to_image

    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"equation_{index:02d}.png"
    try:
        math_to_image(expression, path, dpi=220, format="png", color="black")
        _add_image_before(
            document,
            anchor,
            path,
            width=4.8,
            alt_text=f"公式：{expression.strip('$ ')}",
        )
    except Exception:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(expression.strip("$ "))
        _set_run_font(run, "Cambria Math", 11)
        _move_before(anchor, paragraph._p)


def _insert_blocks(document, anchor, markdown, markdown_path, equation_dir):
    equation_index = 0
    for block in parse_blocks(markdown):
        kind = block[0]
        if kind == "image":
            path = (markdown_path.parent / block[2]).resolve()
            if not path.is_file():
                raise FileNotFoundError(f"report image not found: {path}")
            width = 5.2 if "case_evidence" in path.stem else 5.8
            _add_image_before(document, anchor, path, width=width, alt_text=block[1])
        elif kind == "table":
            rows = block[1]
            if not rows:
                continue
            table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            for row_index, row in enumerate(rows):
                for column_index, value in enumerate(row):
                    table.cell(row_index, column_index).text = value
            _configure_table(table)
            _move_before(anchor, table._tbl)
        elif kind == "equation":
            equation_index += 1
            _add_equation_before(document, anchor, block[1], equation_dir, equation_index)
        else:
            _add_paragraph_before(document, anchor, block[1], kind)


def build_report(template_path, markdown_path, output_path, project_root=None):
    template_path = Path(template_path)
    markdown_path = Path(markdown_path)
    output_path = Path(output_path)
    project_root = Path(project_root) if project_root else markdown_path.resolve().parents[1]
    sections = split_report_sections(markdown_path)
    document = Document(template_path)

    cover_fields = {
        "作品名称": f"作品名称：{PROJECT_TITLE}",
        "作品类型": "作品类型：      开放式自由命题",
        "电子邮箱": "电子邮箱：",
        "提交日期": "提交日期：",
    }
    for paragraph in document.paragraphs:
        for label, value in cover_fields.items():
            if paragraph.text.strip().startswith(label):
                _replace_paragraph(paragraph, value)

    headings, anchors = _section_anchors(document)
    for key in SECTION_HEADINGS:
        _clear_section_content(document, headings[key], anchors[key])
        _insert_blocks(
            document,
            anchors[key],
            sections[key],
            markdown_path,
            output_path.parent / ".report_equations" / key,
        )

    for section in document.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.12)
        section.right_margin = Inches(1.02)

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    _replace_static_total_page_count(document)

    core = document.core_properties
    core.title = PROJECT_TITLE
    core.subject = "2026 第二届大学生人工智能安全竞赛作品报告"
    core.keywords = "TraceGuard, AIGC, fake image detection, explainability"
    core.author = "TraceGuard 参赛团队"
    core.last_modified_by = "TraceGuard 参赛团队"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main(argv=None):
    parser = argparse.ArgumentParser(description="Build the official TraceGuard report DOCX")
    parser.add_argument("--template", required=True)
    parser.add_argument("--source", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args(argv)
    path = build_report(args.template, args.source, args.output)
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
