from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports" / "TraceGuard_作品简介.md"
OUTPUT = ROOT / "作品" / "TraceGuard_作品简介.docx"

NAVY = RGBColor(23, 54, 93)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(95, 95, 95)
FONT = "SimSun"


def set_font(run, size, color=RGBColor(0, 0, 0), bold=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def para(doc, text, *, size=10.5, color=RGBColor(0, 0, 0), bold=False, first=True, before=0, after=6, line=1.15, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.first_line_indent = Cm(0.74) if first else Cm(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size, color=color, bold=bold)
    return p


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.4)
    sec.bottom_margin = Cm(1.4)
    sec.left_margin = Cm(2.25)
    sec.right_margin = Cm(2.25)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("TraceGuard | 作品简介"), 8.5, GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("TraceGuard"), 8.5, GRAY)

    para(doc, "作品简介", size=11.5, color=BLUE, bold=True, first=False, before=2, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "TraceGuard", size=21, color=NAVY, bold=True, first=False, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "TraceGuard:面向社交媒体网络传播的 可解释 AIGC 图像取证平台", size=11, color=GRAY, first=False, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "多模态协同研判 · 真实网络传播 · 超监管分级处置", size=10, color=BLUE, bold=True, first=False, after=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    for line in lines:
        if not line.strip() or line.startswith("# "):
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(line[3:].strip())
            set_font(r, 11, BLUE, True)
            continue
        if line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. ") or line.startswith("4. ") or line.startswith("5. ") or line.startswith("6. "):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.37)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            set_font(p.add_run(line.split(". ", 1)[1]), 9.3)
            continue
        para(doc, line.strip(), size=9.5, after=3, line=1.05)

    doc.core_properties.title = "TraceGuard 作品简介"
    doc.core_properties.subject = "竞赛作品简介"
    doc.core_properties.author = "TraceGuard"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
