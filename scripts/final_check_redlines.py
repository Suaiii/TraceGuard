#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""封版红线终检扫描器（TraceGuard）。

用途：07-20 封版前一键扫描"进入提交包/报告"的内容是否触碰红线，
把此前手工核查（禁词/裸路径/SHA/姓名邮箱/开发痕迹）沉淀成可复跑、可留痕的脚本。

扫描两个层级：
  - 报告层（最严）：reports/TraceGuard.md 源稿 + output/doc/ 的 docx 成品
    （作品报告、原创性声明）。禁一切身份措辞、真实姓名、非中性邮箱、
    裸绝对路径、SHA-256、权重文件名、开发占位词。
  - 程序层：提交包白名单内的源码/配置/前端（镜像 build_submission_package.ps1）。
    禁真实姓名、身份措辞、硬编码绝对路径、ExImage/ExDA 内部资源引用、开发占位；
    SHA/权重名仅在 verified_results/provenance.json 与 REPRODUCIBILITY.md 中豁免。

退出码：有 HIGH 命中 → 2；仅有 REVIEW 命中 → 1；干净 → 0。
结果同时打印并写入 output/final_check_report.txt（UTF-8）。

用法：E:/anaconda/python.exe scripts/final_check_redlines.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]

# ---- 提交包白名单（镜像 build_submission_package.ps1）----
ROOT_ALLOWLIST = {
    "README.md", "REPRODUCIBILITY.md", "batch_analyze.py", "calibrate_risk.py",
    "classify_cases.py", "eval_results.csv", "evaluate_localization.py",
    "requirements-dev.txt", "requirements.txt", "run_test.py", "server.py",
    "start_traceguard.bat", "eval.py", "train.py",
}
DIR_ALLOWLIST = ("configs/", "detection/", "explanation/", "experiments/", "tests/", "web/")
TEXT_SUFFIXES = {".py", ".md", ".txt", ".css", ".js", ".html", ".json",
                 ".yaml", ".yml", ".ps1", ".bat", ".csv", ".cfg", ".ini"}

# 报告源稿 + 成品
REPORT_MD = PROJECT_ROOT / "reports" / "TraceGuard.md"
REPORT_DOCX_DIR = PROJECT_ROOT / "output" / "doc"

# ---- 红线规则 ----
# severity: HIGH（必须清零）| REVIEW（需人工确认，可能是误报）
# scope: report | program | both
# exempt_substr: 命中文件路径含这些子串时豁免该规则
RULES = [
    dict(name="身份措辞", severity="HIGH", scope="both",
         pattern=r"实验室|指导教师|导师|同组|课题组|本校|我校|院系|学院|系主任|师兄|师姐|教研室",
         exempt=["DEVLOG", "AGENTS", "final_check_redlines", "report_restructure"]),
    dict(name="真实姓名", severity="HIGH", scope="report",
         pattern=r"朱羿帅|张潇|贺杰",
         exempt=[]),
    dict(name="真实姓名(程序层-需确认)", severity="REVIEW", scope="program",
         pattern=r"朱羿帅|张潇|贺杰|Suaiii|zx973",
         exempt=["final_check_redlines"]),
    dict(name="邮箱地址", severity="REVIEW", scope="both",
         pattern=r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}",
         exempt=["final_check_redlines"]),
    # 驱动器盘符前不得紧跟字母（否则会把 https:// 里的 s:/ 误判为 D:/ 之类路径）
    dict(name="裸绝对路径", severity="HIGH", scope="report",
         pattern=r"(?<![A-Za-z0-9])[A-Za-z]:[\\/]|/Users/|/home/[a-z]",
         exempt=[]),
    dict(name="硬编码绝对路径(程序层-需确认)", severity="REVIEW", scope="program",
         pattern=r"(?<![A-Za-z0-9])[A-Za-z]:[\\/]|/Users/|/home/[a-z]",
         exempt=["final_check_redlines", "provenance.json"]),
    dict(name="SHA-256", severity="HIGH", scope="report",
         pattern=r"\b[0-9a-fA-F]{64}\b",
         exempt=[]),
    dict(name="SHA-256(程序层)", severity="REVIEW", scope="program",
         pattern=r"\b[0-9a-fA-F]{64}\b",
         exempt=["provenance.json", "verified_results", "REPRODUCIBILITY"]),
    dict(name="权重文件名", severity="HIGH", scope="report",
         pattern=r"best\.pth|ExDA_weights|mambaout_small\.pth|\.ckpt\b",
         exempt=[]),
    # 只抓真正的内部资源；公开的 ExImage.zip（ref[9] 公开数据集）在我们自研复现代码中合法引用，不算违规
    dict(name="ExImage/ExDA内部资源", severity="HIGH", scope="both",
         pattern=r"ExDA_weights|ExDA/|ExDA_|超监管\.pdf|eximage_v2|ExImage-v2|ExImage_v2|ExImage-v2\.\w+",
         exempt=["final_check_redlines", "DEVLOG", "AGENTS"]),
    dict(name="开发占位/未完成痕迹", severity="REVIEW", scope="both",
         pattern=r"\bTODO\b|\bFIXME\b|\bXXX\b|\bHACK\b|待补|待复核|待确认|占位|placeholder|\bTBD\b|【新增】|【待",
         exempt=["final_check_redlines", "DEVLOG", "AGENTS", "test_"]),
]


def program_files() -> list[Path]:
    out = []
    for rel in ROOT_ALLOWLIST:
        p = PROJECT_ROOT / rel
        if p.is_file() and p.suffix.lower() in TEXT_SUFFIXES:
            out.append(p)
    for d in DIR_ALLOWLIST:
        base = PROJECT_ROOT / d.rstrip("/")
        if base.is_dir():
            for p in base.rglob("*"):
                if p.is_file() and p.suffix.lower() in TEXT_SUFFIXES:
                    out.append(p)
    return out


def read_text(p: Path) -> str:
    try:
        return p.read_text(encoding="utf-8", errors="replace")
    except Exception as e:  # pragma: no cover
        return f"<<READ_ERROR {e}>>"


def read_docx_text(p: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return "<<python-docx 未安装，跳过 docx 文本提取>>"
    try:
        d = docx.Document(str(p))
        parts = [para.text for para in d.paragraphs]
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:  # pragma: no cover
        return f"<<DOCX_ERROR {e}>>"


def scan_text(rel_label: str, text: str, scope: str, hits: list):
    lines = text.splitlines()
    for rule in RULES:
        if rule["scope"] not in (scope, "both"):
            continue
        if any(sub in rel_label for sub in rule["exempt"]):
            continue
        rx = re.compile(rule["pattern"])
        for i, line in enumerate(lines, 1):
            for m in rx.finditer(line):
                frag = line.strip()
                if len(frag) > 160:
                    frag = frag[:157] + "..."
                hits.append((rule["severity"], rule["name"], rel_label, i, m.group(0), frag))


def main() -> int:
    hits: list = []

    # ---- 报告层 ----
    report_targets = []
    if REPORT_MD.is_file():
        report_targets.append(("reports/TraceGuard.md", read_text(REPORT_MD)))
    if REPORT_DOCX_DIR.is_dir():
        for docx_path in sorted(REPORT_DOCX_DIR.glob("*.docx")):
            report_targets.append((f"output/doc/{docx_path.name}", read_docx_text(docx_path)))
    for label, text in report_targets:
        scan_text(label, text, "report", hits)

    # ---- 程序层 ----
    for p in program_files():
        rel = p.relative_to(PROJECT_ROOT).as_posix()
        scan_text(rel, read_text(p), "program", hits)

    # ---- 汇总 ----
    highs = [h for h in hits if h[0] == "HIGH"]
    reviews = [h for h in hits if h[0] == "REVIEW"]

    lines_out = []
    lines_out.append("TraceGuard 封版红线终检报告")
    lines_out.append("=" * 60)
    lines_out.append(f"报告层扫描目标：{len(report_targets)} 个（源稿 + docx 成品）")
    lines_out.append(f"程序层扫描文件：{len(program_files())} 个（提交包白名单）")
    lines_out.append(f"HIGH（必须清零）：{len(highs)}    REVIEW（需人工确认）：{len(reviews)}")
    lines_out.append("")

    def emit(title, group):
        lines_out.append(f"--- {title}（{len(group)}）---")
        if not group:
            lines_out.append("  （无）")
            lines_out.append("")
            return
        for sev, name, label, ln, tok, frag in group:
            lines_out.append(f"  [{name}] {label}:{ln}  «{tok}»")
            lines_out.append(f"      {frag}")
        lines_out.append("")

    emit("HIGH 命中", highs)
    emit("REVIEW 命中", reviews)

    if not highs and not reviews:
        lines_out.append("✅ 全部干净，可封版。")
    elif not highs:
        lines_out.append("⚠️ 无 HIGH；REVIEW 项请人工逐条确认（多为合法上下文的误报）。")
    else:
        lines_out.append("❌ 存在 HIGH 命中，封版前必须清零。")

    report_text = "\n".join(lines_out)
    out_path = PROJECT_ROOT / "output" / "final_check_report.txt"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(report_text, encoding="utf-8")
    print(report_text)
    print(f"\n[written] {out_path.relative_to(PROJECT_ROOT).as_posix()}")

    return 2 if highs else (1 if reviews else 0)


if __name__ == "__main__":
    sys.exit(main())
