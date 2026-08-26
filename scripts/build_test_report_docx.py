from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports" / "TraceGuard_测试报告.md"
OUTPUT = ROOT / "作品" / "TraceGuard_测试报告.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"
BODY_FONT = "SimSun"
LATIN_FONT = "Calibri"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=10.5, bold=False, color="000000", italic=False):
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(p, before=0, after=6, line=1.15, first_indent=True):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_indent:
        fmt.first_line_indent = Cm(0.74)


def add_para(doc, text, *, style=None, first_indent=True, before=0, after=6, line=1.15, bold=False, color="000000", size=10.5, italic=False, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    style_paragraph(p, before=before, after=after, line=line, first_indent=first_indent)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=15, bold=True, color=NAVY)
    elif level == 2:
        set_run_font(r, size=12.5, bold=True, color=BLUE)
    else:
        set_run_font(r, size=11, bold=True, color=NAVY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.37)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=10.2)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F7F7F7")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
    r._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("333333")
    return p


def split_table_row(line):
    line = line.strip().strip("|")
    return [x.strip() for x in line.split("|")]


def is_separator(line):
    cells = split_table_row(line)
    return len(cells) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def add_table(doc, rows):
    cols = max(len(row) for row in rows)
    # Content-aware named override: wider description columns and compact numeric columns.
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [2400, 3300, 3660]
    elif cols == 4:
        widths = [1800, 2400, 2400, 2760]
    elif cols == 5:
        widths = [1700, 1500, 1800, 1900, 2460]
    elif cols == 7:
        widths = [1600, 1200, 1250, 1250, 1250, 1400, 1410]
    elif cols == 8:
        widths = [1550, 1150, 1150, 1100, 1100, 1100, 1100, 1110]
    else:
        widths = [9360 // cols] * cols
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    # Keep table headers visible when a table continues on the next page.
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    header_pr.append(header_flag)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, size=8.7, bold=(r_idx == 0), color=(NAVY if r_idx == 0 else "222222"))
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "FBFCFD")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def set_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def configure_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        st = styles[name]
        st.font.name = BODY_FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        st.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("TraceGuard | 测试报告")
    set_run_font(r, size=8.5, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number(footer)


def parse_markdown(doc, lines):
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip("\n")
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code(doc, "\n".join(code_lines))
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line == "---" or line.startswith("报告版本：") or line.startswith("报告日期："):
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            add_bullet(doc, re.sub(r"^\d+\. ", "", line))
            i += 1
            continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                current = lines[i].strip()
                if not is_separator(current):
                    rows.append(split_table_row(current))
                i += 1
            add_table(doc, rows)
            continue
        add_para(doc, line.strip())
        i += 1


def main():
    doc = Document()
    configure_doc(doc)
    # Editorial report cover override: restrained title block, no decorative border.
    add_para(doc, "技术测试与验证报告", first_indent=False, before=12, after=10, size=12, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "TraceGuard", first_indent=False, before=0, after=4, size=25, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "TraceGuard:面向社交媒体网络传播的 可解释 AIGC 图像取证平台", first_indent=False, before=0, after=22, size=13, color="444444", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "测试对象：检测器、解释与定位模块、风险融合、Web/API/CLI 运行链路", first_indent=False, after=5, size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "报告版本：v1.1　　报告日期：2026 年 8 月 26 日", first_indent=False, after=20, size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF5FB")
    pPr.append(shd)
    r = p.add_run("核心结论：历史冻结基线 191/191 通过；当前工作区全量复测 216 项中 213 项通过，3 项为批量 CLI 超时或提交包装目录缺失。平台 Web/API/CLI 闭环可运行。跨域、传播、定位和风险测试均已保留数据边界，局部框当前仅作为审核线索，不作为像素级篡改定论。")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    doc.add_page_break()
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8").splitlines())
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "TraceGuard 测试报告"
    doc.core_properties.subject = "TraceGuard 技术测试与验证"
    doc.core_properties.author = "TraceGuard"
    doc.core_properties.comments = "Generated from reports/TraceGuard_测试报告.md"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
